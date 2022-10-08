from datetime import datetime
import os
import re
import shutil
from urllib.request import urlopen, Request
import requests
import aspose.words as aw
from bs4 import BeautifulSoup
from docx import Document
from py7zr import unpack_7zarchive
from pyunpack import Archive
import pandas as pd
from win32com import client as wc

hdr = {'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) '
                     'Chrome/55.0.2883.87 Safari/537.36'}

keywords_samples = ['товарные образцы',
                    'экземпляр СИЗ',
                    'экземпляр продукции',
                    'экземпляр материала',
                    'образец направленный',
                    'с даты заключения договора опытный образец',
                    'утвержденным образцам',
                    'опытный образец',
                    'товарный образец',
                    'образец товара',
                    'образцы товаров',
                    'экземпляры СИЗ',
                    'экземпляр продукта',
                    'образцы направленные',
                    'опытные образцы',
                    'утвержденный образец',
                    'предоставить образцы',
                    'предоставить образец',
                    'предоставлять образцы',
                    'предоставлять образец',
                    'образец для опытной ',
                    'образцы для опытной ',
                    'согласования образцов',
                    'согласование образца',
                    'согласование образцов',
                    'направить образец',
                    'направить образцы',
                    'направленные образцы']

keywords_analogs = ['аналог',
                    'аналоги',
                    'аналогичной',
                    'аналогичной продукции',
                    'аналогичного товара',
                    'аналогичных товаров',
                    'эквивалент',
                    'эквивалента',
                    'эквиваленты',
                    'аналогичный']

keywords_delivery_time = ['срок поставки',
                          'сроки поставки',
                          'сроки поставок',
                          'срок выполнения поставок',
                          'срок выполнения поставки',
                          'сроки выполнения поставок',
                          'срок завершения поставок',
                          'срок завершения поставки',
                          'сроки завершения поставок',
                          'сроки завершения поставки',
                          'срок поставки товара',
                          'поставка товара осуществляется',
                          'поставка товаров осуществляется',
                          'график поставки',
                          'графики поставки',
                          'поставка партии',
                          'поставки партии',
                          'должна быть поставлена',
                          'сроки поставки товара',
                          'поставка продукции осуществляется',
                          'поставка товара',
                          'поставка товаров',
                          'график поставок',
                          'график выполнения поставки',
                          'график выполнения поставок',
                          'поставка партии',
                          'поставки партии',
                          'сроки (периоды)']

keywords_payment_time = [
    'срок оплаты',
    'сроки оплаты',
    'условия оплаты',
    'порядок оплаты',
    'платеж в размере',
    'оплата поставленной',
    'оплата продукции',
    'оплата поставленного',
    'оплата продукции',
    'оплата товара',
    'оплата за приобретаемый',
    'оплата за приобретаемую',
    'оплата приобретенного',
    'оплата приобретенной',
    'оплата за приобретенный',
    'оплата за приобретенную',
    'расчет за поставленный',
    'расчет за поставленную',
    'оплачивает поставленный',
    'оплачивает поставленную',
    'оплата за товар',
    'оплата за продукцию',
    'производит оплату',
    'оплата в течение',
    'оплачивает приобретенный',
    'порядок расчетов',
    'расчеты по договору',
    'расчет по договору',
    'производит расчет'
]

keywords_divisibility = ['количество лотов',
                         'попозиционная поставка',
                         '2х и более победителей',
                         'нескольких победителей',
                         'выбор более одного победителя',
                         'несколько поставщиков']

keywords_address = ['адрес поставки',
                    'адреса поставки',
                    'адреса поставок',
                    'место поставки',
                    'место и порядок',
                    'место доставки',
                    'доставка до склада',
                    'должен быть доставлен',
                    'адрес склада',
                    'место доставки',
                    'адрес грузополучателя',
                    'склад покупателя',
                    'поставка товара',
                    'адрес доставки',
                    'адреса доставок',
                    'адреса доставки',
                    'адрес']

keywords_support = ['обеспечение договора',
                    'обеспечение заявки',
                    'обеспечение исполнения',
                    'обеспечение исполнения обязательств',
                    'размер обеспечения',
                    'банковское сопровождение',
                    'обеспечение обязательств']

procedure_number = []
customer = []
method_of_conducting = []
date_of_placement = []
end_date = []
nmc = []
electronic_platform = []
data_samples = []
data_analogs = []
data_delivery_time = []
data_payment_time = []
data_divisibility = []
data_address = []
data_support = []
data_url = []
data_dict = {'Номер процедуры': procedure_number,
             'Заказчик': customer,
             'Способ проведения': method_of_conducting,
             'Дата размещения': date_of_placement,
             'Дата окончания подачи заявок': end_date,
             'НМЦ': nmc,
             'Электронная площадка': electronic_platform,
             'Образцы': data_samples,
             'Аналоги': data_analogs,
             'Срок поставки': data_delivery_time,
             'Срок оплаты': data_payment_time,
             'Делимость': data_divisibility,
             'Адрес': data_address,
             'Обеспечение': data_support,
             'Ссылка': data_url}


def main():
    start_time = datetime.now()
    file1 = open("urls.txt", "r")
    while True:
        # считываем строку
        line = file1.readline()
        # прерываем цикл, если строка пустая
        if not line:
            break

        if os.path.isdir("content"):
            shutil.rmtree('content')
            os.mkdir("content")
        else:
            os.mkdir("content")
        if 'zakupki.gov.ru/epz/order/notice' in line:
            parse_epz_order(line)
        if 'zakupki.gov.ru/223/purchase/public' in line:
            parse_223_order(line)
        if 'zakupki.gov.ru/epz/pricereq/card' in line:
            parse_order_card(line)

    # закрываем файл
    file1.close()
    write_dataframe_to_excel(data_dict)
    print(datetime.now() - start_time)


def parse_order_card(line):
    while True:
        try:
            procedure_str = ''
            customer_str = ''
            method_str = ''
            date_str = ''
            end_date_str = ''
            platform_str = ''
            nmc_str = ''
            url = line.strip()
            r = requests.get(url, headers=hdr)
            html_page = urlopen(url)
            soup = BeautifulSoup(html_page, features="lxml")
            # div_1 = soup.find("div", {"class": "span-body"})  # Parse the first div.
            # div_2 = div_1("div", {"class": "timestamp updated"})
            number = soup.find("span", {"class": "navBreadcrumb__text"})
            procedure_str += number.text.strip()
            # span = soup.find('span', {'class': "cardMainInfo__purchaseLink distancedText"})
            # print(span.text.strip())
            divs = soup.findAll('div', {'class': "cardMainInfo__section col-6"})
            for div in divs:
                spans = div.findAll('span')
                if spans[0].text.strip() == 'Размещено':
                    date_str += spans[1].text.strip()

            divs = soup.findAll('div', {'class': "cardMainInfo__section col-12"})
            for div in divs:
                spans = div.findAll('span')
                if spans[0].text.strip() == 'Сроки проведения закупки':
                    end_date_str += spans[0].text.strip() + ': ' + spans[1].text.strip()

            divs = soup.findAll('div', {'class': 'cardMainInfo__section'})
            for div in divs:
                spans = div.findAll('span')
                if spans[0].text.strip() == 'Организация, разместившая запрос цен':
                    customer_str += spans[0].text.strip() + ': ' + spans[1].text.strip()
            method_str += 'Процедура сбора информации'
            nmc_str += 'Не указано'
            platform_str += 'Не указано'
            add_parse_str_to_list(customer_str, date_str, end_date_str, method_str, nmc_str, platform_str,
                                  procedure_str)
            break
        except Exception:
            pass
    url_epz = line.replace('common-info.html', 'docs.html')
    while True:
        try:
            req = Request(
                url_epz,
                data=None,
                headers={
                    'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_9_3) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/35.0.1916.47 Safari/537.36'
                }
            )

            f = urlopen(req)
            html_page = f.read().decode('utf-8')
            soup = BeautifulSoup(html_page, features="lxml")
            for link_data in soup.findAll('div', {'class': 'price_file'}):
                link = link_data.findAll('a')[1]
                filename = link.text.strip()
                url_download = link.get('href')
                print(filename)
                print(url_download)
                while True:
                    try:
                        file_object = requests.get(url_download, headers=hdr)
                        filename = check_file_extension(file_object, filename)
                        with open(f'content/{filename}', 'wb') as local_file:
                            local_file.write(file_object.content)
                            local_file.close()
                            extract_files_from_archive(filename)
                        break
                    except Exception:
                        pass
            break
        except Exception as e:
            print(e)
            pass
    parse_docs_from_dir(url_epz)


def parse_223_order(line):
    doc_url, r, r_docs, url = parse_url(line.strip())
    tries = 0
    need_to_parse = True
    while True:
        try:
            html_page = get_html_page(doc_url)
            soup = BeautifulSoup(html_page, "lxml")
            need_to_parse = download_docs(soup, url)
            break
        except Exception:
            if tries >= 30:
                write_error_to_file(url)
                break
            tries += 1
            pass
    if need_to_parse:
        parse_docs_from_dir(url)


def parse_epz_order(line):
    parse_url_epz(line)
    url_epz = line.replace('common-info.html', 'documents.html')
    print(url_epz)
    download_docs_epz(url_epz)
    parse_docs_from_dir(url_epz)


def download_docs_epz(url_epz):
    while True:
        try:
            req = Request(
                url_epz,
                data=None,
                headers={
                    'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_9_3) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/35.0.1916.47 Safari/537.36'
                }
            )

            f = urlopen(req)
            html_page = f.read().decode('utf-8')
            soup = BeautifulSoup(html_page, features="lxml")
            for link_data in soup.findAll('div', {'class': 'col clipText'}):
                link = link_data.findAll('a')[1]
                filename = link.text.strip()
                url_download = link.get('href')
                print(filename)
                print(url_download)
                while True:
                    try:
                        file_object = requests.get(url_download, headers=hdr)
                        filename = check_file_extension(file_object, filename)
                        with open(f'content/{filename}', 'wb') as local_file:
                            local_file.write(file_object.content)
                            local_file.close()
                            extract_files_from_archive(filename)
                        break
                    except Exception:
                        pass
            break
        except Exception as e:
            print(e)
            pass


def parse_docs_from_dir(url):
    check_dirs = check_dirs_in_content()
    while check_dirs:
        check_dirs = check_dirs_in_content()
    # fix fix fix       doc to docx
    # paths = []
    # folder = os.getcwd()
    # for root, dirs, files in os.walk(folder):
    #     for file in files:
    #         if file.endswith('doc') and not file.startswith('~'):
    #             paths.append(os.path.join(root, file))
    # for path in paths:
    #     w = wc.Dispatch('Word.Application')
    #     doc = w.Documents.Open(path)
    #     doc.SaveAs(path + "x", 16)
    #     doc.Close()
    # w.Quit()
    parse_docs(url)
    shutil.rmtree('content')


def parse_url_epz(line):
    while True:
        try:
            procedure_str = ''
            customer_str = ''
            method_str = ''
            date_str = ''
            end_date_str = ''
            platform_str = ''
            nmc_str = ''
            url = line.strip()
            r = requests.get(url, headers=hdr)
            html_page = urlopen(url)
            soup = BeautifulSoup(html_page, features="lxml")
            # div_1 = soup.find("div", {"class": "span-body"})  # Parse the first div.
            # div_2 = div_1("div", {"class": "timestamp updated"})
            number = soup.find("span", {"class": "navBreadcrumb__text"})
            procedure_str += number.text.strip()
            print(procedure_str)
            divs = soup.findAll('div', {"class": "cardMainInfo__section"})
            for div in divs:
                spans = div.findAll('span')
                if spans[0].text.strip() == 'Заказчик' or spans[
                    0].text.strip() == 'Организация, осуществляющая размещение':
                    customer_str += spans[1].text.strip()
                if spans[0].text.strip() == 'Размещено':
                    date_str += spans[1].text.strip()
                if spans[0].text.strip() == 'Окончание подачи заявок':
                    end_date_str += spans[1].text.strip()
            divs = soup.find('div', {"class": "price"})
            spans = divs.findAll('span')
            if spans[0].text.strip() == 'Начальная цена':
                nmc_str += spans[1].text.strip()
            sections = soup.findAll('section', {'class': "blockInfo__section section"})
            for section in sections:
                spans = section.findAll('span')
                if 'Адрес электронной площадки в информационно-телекоммуникационной сети' in spans[0].text.strip():
                    platform_str += spans[1].text.strip()
            method_str += 'нет'
            add_parse_str_to_list(customer_str, date_str, end_date_str, method_str, nmc_str, platform_str,
                                  procedure_str)
            break
        except Exception:
            pass


def write_error_to_file(url):
    my_file = open("logs.txt", "a+")
    logs = f'{url}\t is not correct\n'
    my_file.write(logs)
    my_file.close()


def write_dataframe_to_excel(data_to_excel):
    df = pd.DataFrame(data_to_excel)
    if os.path.isfile('data.xlsx'):
        old_df = pd.read_excel('data.xlsx', index_col=0)
        df = old_df.append(df, ignore_index=True)

    df.to_excel("data.xlsx")


def parse_docs(url):
    str_samples = ''
    str_analogs = ''
    str_delivery_time = ''
    str_payment_time = ''
    str_divisibility = ''
    str_address = ''
    str_support = ''

    for filename in os.listdir("content"):
        print(filename)
        split_tup = os.path.splitext(filename)
        extension = split_tup[1].lower()
        try:
            str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples, str_support \
                = get_info_from_docx(extension, filename, str_address, str_analogs, str_delivery_time, str_divisibility,
                                     str_payment_time, str_samples, str_support)
        except Exception:
            pass

    add_data_to_list(str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples,
                     str_support, url)


def get_info_from_docx(extension, filename, str_address, str_analogs, str_delivery_time, str_divisibility,
                       str_payment_time, str_samples, str_support):
    if '.pdf' == extension:
        doc = aw.Document(f'content/{filename}')
        print(filename + '\n\n')
        split_tup = os.path.splitext(filename)
        extension = '.docx'
        filename = split_tup[0] + extension
        print(filename)
        doc.save(f'content/{filename}')
    if '.docx' == extension:
        try:
            doc = Document(f'content/{filename}')
        except Exception:
            return
        check_head = []
        for para in doc.paragraphs:
            try:
                str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples, str_support = find_keywords(
                    para, str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples,
                    str_support, filename)
                if "Heading" not in para.style.name:
                    run_bold_text = ''
                    for run in para.runs:
                        if run.bold:
                            run_bold_text = run_bold_text + ' ' + run.text
                            run_bold_text = run_bold_text.strip()
                    # print(run_bold_text)
                    if run_bold_text != '':
                        check_head = check_content_in_headers(run_bold_text)
                        # print(check_head)
                    if check_head and check_head[0]:
                        if check_head[1] == 'samples':
                            if filename not in str_samples:
                                str_samples += filename.strip() + ':\n'
                            str_samples += '\n' + para.text.strip()
                        if check_head[1] == 'analogs':
                            if filename not in str_analogs:
                                str_analogs += filename.strip() + ':\n'
                            str_analogs += '\n' + para.text.strip()
                        if check_head[1] == 'delivery_time':
                            if filename not in str_delivery_time:
                                str_delivery_time += filename.strip() + ':\n'
                            str_delivery_time += '\n' + para.text.strip()
                        if check_head[1] == 'payment_time':
                            if filename not in str_payment_time:
                                str_payment_time += filename.strip() + ':\n'
                            str_payment_time += '\n' + para.text.strip()
                        if check_head[1] == 'address':
                            if filename not in str_address:
                                str_address += filename.strip() + ':\n'
                            str_address += '\n' + para.text.strip()
                        if check_head[1] == 'divisibility':
                            if filename not in str_divisibility:
                                str_divisibility += filename.strip() + ':\n'
                            str_divisibility += '\n' + para.text.strip()
                        if check_head[1] == 'support':
                            if filename not in str_support:
                                str_support += filename.strip() + ':\n'
                            str_support += '\n' + para.text.strip()
            except Exception as e:
                print(e)
        for table in doc.tables:
            try:
                for row in table.rows:
                    for cell in row.cells:
                        for key_word in keywords_samples:
                            str_samples = check_length_text(cell, key_word, row, str_samples, filename)
                        for key_word in keywords_analogs:
                            str_analogs = check_length_text(cell, key_word, row, str_analogs, filename)
                        for key_word in keywords_delivery_time:
                            str_delivery_time = check_length_text(cell, key_word, row, str_delivery_time, filename)
                        for key_word in keywords_payment_time:
                            str_payment_time = check_length_text(cell, key_word, row, str_payment_time, filename)
                        for key_word in keywords_address:
                            str_address = check_length_text(cell, key_word, row, str_address, filename)
                        for key_word in keywords_divisibility:
                            str_divisibility = check_length_text(cell, key_word, row, str_divisibility, filename)
                        for key_word in keywords_support:
                            str_support = check_length_text(cell, key_word, row, str_support, filename)

                # Data will be a list of rows represented as dictionaries
                # containing each row's data.
                keys = None
                for i, row in enumerate(table.rows):
                    text = (cell.text.replace('\n', ' ') for cell in row.cells)
                    # Establish the mapping based on the first row
                    # headers; these will become the keys of our dictionary
                    if i == 0:
                        keys = tuple(text)
                        continue
                    # Construct a dictionary for this row, mapping
                    # keys to values for this row
                    row_data = dict(zip(keys, text))
                    for data in row_data:
                        str_delivery_time = get_vertical_info_from_table(data, row_data, str_delivery_time,
                                                                         keywords_delivery_time, filename)
                        str_samples = get_vertical_info_from_table(data, row_data, str_samples,
                                                                   keywords_samples, filename)
                        str_analogs = get_vertical_info_from_table(data, row_data, str_analogs,
                                                                   keywords_analogs, filename)
                        str_payment_time = get_vertical_info_from_table(data, row_data, str_payment_time,
                                                                        keywords_payment_time, filename)
                        str_divisibility = get_vertical_info_from_table(data, row_data, str_divisibility,
                                                                        keywords_divisibility, filename)
                        str_support = get_vertical_info_from_table(data, row_data, str_support,
                                                                   keywords_support, filename)
                        str_address = get_vertical_info_from_table(data, row_data, str_address,
                                                                   keywords_address, filename)
            except Exception as e:
                print(e)
    return str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples, str_support


def get_vertical_info_from_table(data, row_data, str_param, keywords, filename):
    for key_word in keywords:
        if key_word in data:
            if data + ': ' + row_data[data] not in str_param:
                if filename not in str_param:
                    str_param += filename.strip() + ':\n'
                str_param += data + ': ' + row_data[data] + '\n'
    return str_param


def check_length_text(cell, key_word, row, str_param, filename):
    temp = ''
    if key_word in cell.text.lower():
        for cell_temp in row.cells:
            temp += '\n' + cell_temp.text.strip()
    if 300 > len(temp) > 0:
        if temp not in str_param:
            if filename not in str_param:
                str_param += filename.strip() + ':\n'
            str_param += '\n' + temp + '\n'
    return str_param


def check_content_in_headers(run_bold_text):
    res = []
    # образцы
    for key_word in keywords_samples:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('samples')
            return res
    # аналоги
    for key_word in keywords_analogs:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('analogs')
            return res
    # срок доставки
    for key_word in keywords_delivery_time:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('delivery_time')
            return res
    # срок оплаты
    for key_word in keywords_payment_time:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('payment_time')
            return res
    # адрес
    for key_word in keywords_address:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('address')
            return res
    # делимость
    for key_word in keywords_divisibility:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('divisibility')
            return res
    # обеспечение
    for key_word in keywords_support:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('support')
            return res
    if not res:
        res.append(False)
        res.append('none')
    return res


def find_keywords(para, str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples,
                  str_support, filename):
    try:
        if para.text or not para.text.isspace():
            # образцы
            str_samples = get_paragraphs(para, str_samples, keywords_samples, filename)
            # аналоги
            str_analogs = get_paragraphs(para, str_analogs, keywords_analogs, filename)
            # срок доставки
            str_delivery_time = get_paragraphs(para, str_delivery_time, keywords_delivery_time, filename)
            # срок оплаты
            str_payment_time = get_paragraphs(para, str_payment_time, keywords_payment_time, filename)
            # адрес
            str_address = get_paragraphs(para, str_address, keywords_address, filename)
            # делимость
            str_divisibility = get_paragraphs(para, str_divisibility, keywords_divisibility, filename)
            # обеспечение
            str_support = get_paragraphs(para, str_support, keywords_support, filename)
    except Exception as e:
        print(e)
    return str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples, str_support


def get_paragraphs(para, param_str, keywords_list, filename):
    for key_word in keywords_list:
        if key_word in para.text.lower():
            if filename not in param_str:
                param_str += filename.strip() + ':\n'
            if para.text + '\n' not in param_str:
                param_str += para.text + '\n'
    return param_str


def add_data_to_list(str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples,
                     str_support, url):
    if str_samples == '':
        data_samples.append('нет')
    else:
        data_samples.append(str_samples)
    if str_analogs == '':
        data_analogs.append('нет')
    else:
        data_analogs.append(str_analogs)
    if str_delivery_time == '':
        data_delivery_time.append('нет')
    else:
        data_delivery_time.append(str_delivery_time)
    if str_payment_time == '':
        data_payment_time.append('нет')
    else:
        data_payment_time.append(str_payment_time)
    if str_address == '':
        data_address.append('нет')
    else:
        data_address.append(str_address)
    if str_divisibility == '':
        data_divisibility.append('нет')
    else:
        data_divisibility.append(str_divisibility)
    if str_support == '':
        data_support.append('нет')
    else:
        data_support.append(str_support)

    data_url.append(url)


def check_dirs_in_content():
    check_dirs = False
    list_dir = os.listdir("content")
    for dir_temp in list_dir:
        path = 'content/'
        path += dir_temp
        if os.path.isdir(path):
            check_dirs = True
            allfiles = os.listdir(path)
            for f in allfiles:
                os.rename(path + '/' + f, 'content/' + f)
            os.removedirs(path)
        if '.zip' in dir_temp or '.rar' in dir_temp or '.7z' in dir_temp:
            extract_files_from_archive(dir_temp)
            check_dirs = True
    return check_dirs


def download_docs(soup, url):
    need_to_parse = True
    allowed_downloads, count_downloads = check_validity_docs(soup)
    if allowed_downloads == 0:
        need_to_parse = False
        write_error_to_file(url)
    for link in soup.findAll('a', {'class': 'epz_aware'}):
        filename = link.text
        filename = filename.replace('\n', '')
        while True:
            try:
                url_download = "https://zakupki.gov.ru" + link.get('href')
                file_object = requests.get(url_download, headers=hdr)
                filename = check_file_extension(file_object, filename)
                with open(f'content/{filename}', 'wb') as local_file:
                    local_file.write(file_object.content)
                    local_file.close()
                    extract_files_from_archive(filename)
                break
            except Exception as e:
                print(e)
        count_downloads += 1
        if count_downloads >= allowed_downloads:
            break
    return need_to_parse


def check_file_extension(file_object, filename):
    content_type = file_object.headers['Content-Disposition']
    filename_redirect = re.findall("filename=(.+)", content_type)[0]
    split_tup = os.path.splitext(filename_redirect.replace('"', ''))
    extension = split_tup[1]
    if extension not in filename:
        filename += extension
    return filename


def check_validity_docs(soup):
    count_downloads = 0
    allowed_downloads = 0
    for numb in soup.findAll('td', {'style': 'width: 15%'}):
        if "(недействующая)" in numb.text:
            break
        allowed_downloads += 1
    return allowed_downloads, count_downloads


def extract_files_from_archive(filename):
    try:
        if ".rar" in filename or ".zip" in filename:
            Archive(f'content/{filename}').extractall('content')
            os.remove(f'content/{filename}')
        if ".7z" in filename:
            shutil.register_unpack_format('7zip', ['.7z'], unpack_7zarchive)
            shutil.unpack_archive(f'content/{filename}', 'content')
            os.remove(f'content/{filename}')
    except Exception:
        os.remove(f'content/{filename}')


def get_html_page(doc_url):
    count_connections = 0
    while True:
        try:
            count_connections += 1
            if count_connections > 50:
                break
            html_page = urlopen(doc_url)
            return html_page
        except Exception:
            pass


def parse_url(curr_url):
    while True:
        try:
            procedure_str = ''
            customer_str = ''
            method_str = ''
            date_str = ''
            end_date_str = ''
            platform_str = ''
            nmc_str = ''

            url = curr_url
            r = requests.get(url, headers=hdr)
            html_page = urlopen(url)
            soup = BeautifulSoup(html_page, features="lxml")
            rows = soup.findAll('tr')
            for row in rows:
                if row.find('span'):
                    list_text = row.text.strip().split('\n')
                    res_text = []
                    for text in list_text:
                        if not text.isspace() and text:
                            res_text.append(text.strip())
                    if 'Реестровый номер извещения' in res_text[0]:
                        procedure_str += res_text[1] + '\n'
                    if 'Наименование организации' in res_text[0]:
                        customer_str += res_text[1] + '\n'
                    if 'Способ размещения закупки' in res_text[0]:
                        method_str += res_text[1] + '\n'
                    if 'Дата размещения извещения' in res_text[0]:
                        date_str += res_text[1] + '\n'
                    if 'Дата и время окончания подачи заявок' in res_text[0]:
                        end_date_str += res_text[2] + '\n'
                    if 'Адрес электронной площадки' in res_text[0]:
                        platform_str += res_text[1] + '\n'

                if row.find('td'):
                    list_text = row.text.strip().split('\n')
                    res_text = []
                    for text in list_text:
                        if not text.isspace() and text:
                            res_text.append(text.strip())
                    if 'Наименование организации' in row.text:
                        customer_str += res_text[1] + '\n'
                    if 'Дата размещения извещения' in row.text:
                        date_str += res_text[1] + '\n'
                    if 'Адрес электронной площадки' in row.text:
                        platform_str += res_text[1] + '\n'
            break
        except Exception:
            pass
    while True:
        try:
            nmc_url = url.replace("common-info.html", "lot-list.html")
            r = requests.get(nmc_url, headers=hdr)
            html_page = get_html_page(nmc_url)
            soup = BeautifulSoup(html_page, features="lxml")
            rows = soup.findAll('td')
            for row in rows:
                if 'Российский рубль' in row.text:
                    nmc_str += row.text.strip() + '\n'
            break
        except Exception:
            pass

    add_parse_str_to_list(customer_str, date_str, end_date_str, method_str, nmc_str, platform_str, procedure_str)
    doc_url = url.replace("common-info.html", "documents.html")
    r_docs = requests.get(doc_url, headers=hdr)
    return doc_url, r, r_docs, url


def add_parse_str_to_list(customer_str, date_str, end_date_str, method_str, nmc_str, platform_str, procedure_str):
    if not procedure_str:
        procedure_number.append('нет')
    else:
        procedure_number.append(procedure_str)
    if not customer_str:
        customer.append('нет')
    else:
        customer.append(customer_str)
    if not method_str:
        method_of_conducting.append('нет')
    else:
        method_of_conducting.append(method_str)
    if not date_str:
        date_of_placement.append('нет')
    else:
        date_of_placement.append(date_str)
    if not end_date_str:
        end_date.append('нет')
    else:
        end_date.append(end_date_str)
    if not platform_str:
        electronic_platform.append('нет')
    else:
        electronic_platform.append(platform_str)
    if not nmc_str:
        nmc.append('нет')
    else:
        nmc.append(nmc_str)


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    main()
